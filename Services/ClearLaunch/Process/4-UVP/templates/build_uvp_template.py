"""
Build ClearLaunch Step 4 (UVP) .docx workshop template
with ClearThink brand styling. Industry-agnostic version.
"""
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Brand colors
GREEN = RGBColor(0x1B, 0x9B, 0x5E)
BRIGHT = RGBColor(0x3B, 0xEB, 0x96)
DARK = RGBColor(0x12, 0x17, 0x18)
CREAM_HEX = "F6F3EF"
LGGREEN_HEX = "E5F5EC"
ALTROW_HEX = "F0F9F4"
GREEN_HEX = "1B9B5E"
DARK_HEX = "121718"
BORDER_HEX = "A8D9BD"
WHITE_HEX = "FFFFFF"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for side, color in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if color:
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            )
            borders.append(border)
    tcPr.append(borders)


def add_styled_paragraph(doc, text, font_name="Inter", font_size=11, bold=False,
                         color=DARK, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    return p


def add_section_header(doc, text):
    add_styled_paragraph(doc, text, font_size=14, bold=True, color=GREEN, space_after=4)


def add_description(doc, text):
    add_styled_paragraph(doc, text, font_size=9.5, color=RGBColor(0x66, 0x66, 0x66),
                         space_after=2)


def add_adaptive_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"Agent Note: {text}")
    run.font.name = "Inter"
    run.font.size = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1B, 0x9B, 0x5E)


def add_question_table(doc, questions, header_text="Questions"):
    """Create a branded question/answer table."""
    table = doc.add_table(rows=1 + len(questions), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(4.0)
        row.cells[1].width = Inches(2.5)

    # Header row
    hdr = table.rows[0]
    set_cell_shading(hdr.cells[0], GREEN_HEX)
    set_cell_shading(hdr.cells[1], GREEN_HEX)

    for i, text in enumerate([header_text, "Answer"]):
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = "Inter"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Question rows with alternating shading
    for idx, question in enumerate(questions):
        row = table.rows[idx + 1]
        bg = ALTROW_HEX if idx % 2 == 0 else WHITE_HEX

        # Question cell
        set_cell_shading(row.cells[0], bg)
        p = row.cells[0].paragraphs[0]
        run = p.add_run(question)
        run.font.name = "Inter"
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK

        # Answer cell (empty)
        set_cell_shading(row.cells[1], bg)

    # Add borders to all cells
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=BORDER_HEX, bottom=BORDER_HEX,
                           left=BORDER_HEX, right=BORDER_HEX)

    doc.add_paragraph()  # spacer
    return table


def add_client_info_table(doc):
    """Add the client info header table."""
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_fields = [("Client / Business Name", ""), ("Industry / Vertical", ""), ("Date", "")]
    for i, (label, _) in enumerate(info_fields):
        row = table.rows[i]
        set_cell_shading(row.cells[0], LGGREEN_HEX)
        set_cell_border(row.cells[0], top=BORDER_HEX, bottom=BORDER_HEX,
                       left=BORDER_HEX, right=BORDER_HEX)
        set_cell_border(row.cells[1], top=BORDER_HEX, bottom=BORDER_HEX,
                       left=BORDER_HEX, right=BORDER_HEX)

        p = row.cells[0].paragraphs[0]
        run = p.add_run(label)
        run.font.name = "Inter"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = DARK

        p2 = row.cells[1].paragraphs[0]
        run2 = p2.add_run("[________________]")
        run2.font.name = "Inter"
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()


# ========================================
# BUILD UVP TEMPLATE (Step 4)
# ========================================
def build_uvp_template():
    doc = Document()

    # Title
    add_styled_paragraph(doc, "", space_after=12)
    add_styled_paragraph(doc, "UNIQUE VALUE PROPOSITION",
                         font_size=24, bold=True, color=GREEN,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, "Discovery Workshop Template",
                         font_size=14, bold=False, color=DARK,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, "ClearLaunch System | Step 4: UVP Development",
                         font_size=10, color=RGBColor(0x99, 0x99, 0x99),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    # Purpose
    add_styled_paragraph(
        doc,
        "Purpose: This template guides a structured discovery process to identify and articulate "
        "what truly differentiates the business in their marketplace. Use these questions to uncover "
        "meaningful competitive advantages that will form the foundation of their GTM strategy.",
        font_size=10, color=RGBColor(0x66, 0x66, 0x66), space_after=4
    )

    add_styled_paragraph(doc, "Key Guidelines:", font_size=10, bold=True, space_after=2)
    guidelines = [
        "Push for specific, tangible differentiators rather than generic claims",
        "Look for the intersection of what they do exceptionally well and what their market values most",
        "Document examples that demonstrate their UVP in action",
        "Focus on benefits to customers, not just features of their offering",
    ]
    for g in guidelines:
        add_styled_paragraph(doc, f"  \u2022  {g}", font_size=9.5,
                             color=RGBColor(0x66, 0x66, 0x66), space_after=2)

    add_styled_paragraph(
        doc,
        "After completing this exercise, synthesize the responses into a concise UVP statement "
        "(1-2 sentences) that clearly articulates their unique market position.",
        font_size=10, color=RGBColor(0x66, 0x66, 0x66), space_after=12
    )

    # Client info
    add_client_info_table(doc)

    # --- Section 1: Problem & Market Gap ---
    add_section_header(doc, "1. Problem & Market Gap")
    add_adaptive_note(doc,
        "For product businesses, focus on what gap the product fills. For services, focus on what "
        "outcome they deliver that others miss. For SaaS, probe the workflow or process pain the "
        "software eliminates.")
    add_question_table(doc, [
        "What specific problem does your business solve that others in your space don't address effectively?",
        "What gaps or underserved needs in your market does your business specifically address?",
        "What specific pain points do you eliminate for customers that others typically overlook?",
    ], "Problem & Market Gap")

    # --- Section 2: Methodology & Approach ---
    add_section_header(doc, "2. Methodology & Approach")
    add_adaptive_note(doc,
        "For services, probe proprietary frameworks or delivery methods. For products, probe design "
        "philosophy, sourcing, or manufacturing. For SaaS, probe the technical architecture or UX "
        "philosophy that drives the product.")
    add_question_table(doc, [
        "What unique methodology, process, or approach sets your offerings apart from competitors?",
        "How do you deliver differently than competitors (faster, more transparent, more hands-on, more automated, etc.)?",
        "What unique combination of products, services, or capabilities do you offer that provides a comprehensive solution?",
    ], "Methodology & Approach")

    # --- Section 3: Expertise & Credibility ---
    add_section_header(doc, "3. Expertise & Credibility")
    add_adaptive_note(doc,
        "For product companies, probe sourcing expertise, R&D investment, or proprietary technology. "
        "For services, probe certifications, years in niche, or specialized training. For founders, "
        "probe personal story and domain authority.")
    add_question_table(doc, [
        "What specialized expertise, certifications, technology, or experience does your team possess that's rare in your space?",
        "What aspects of your company culture or values translate into tangible benefits for customers?",
    ], "Expertise & Credibility")

    # --- Section 4: Results & Outcomes ---
    add_section_header(doc, "4. Results & Outcomes")
    add_adaptive_note(doc,
        "For B2B, focus on measurable business outcomes (revenue, efficiency, cost savings). For "
        "B2C/DTC, focus on customer experience outcomes (satisfaction, lifestyle improvement, time "
        "saved). For SaaS, probe specific metrics the platform improves.")
    add_question_table(doc, [
        "What specific results or outcomes can customers expect that they can't get elsewhere?",
        "What specific metrics or KPIs can you improve for customers better than competitors?",
    ], "Results & Outcomes")

    # --- Section 5: Risk Reduction & Assurance ---
    add_section_header(doc, "5. Risk Reduction & Assurance")
    add_adaptive_note(doc,
        "For high-ticket services, probe guarantees and contracts. For ecommerce/DTC, probe return "
        "policies, trials, and social proof mechanisms. For SaaS, probe free tiers, onboarding "
        "support, and data portability.")
    add_question_table(doc, [
        "What guarantees, assurances, or risk-reduction elements can you offer that competitors don't?",
        "How does your pricing model or structure provide better value than traditional approaches in your space?",
    ], "Risk Reduction & Assurance")

    # --- Section 6: Proof & Validation ---
    add_section_header(doc, "6. Proof & Validation")
    add_adaptive_note(doc,
        "For B2B, focus on case studies and named client outcomes. For B2C/DTC, focus on reviews, "
        "user-generated content, and community evidence. For SaaS, probe usage metrics, retention "
        "rates, and customer stories.")
    add_question_table(doc, [
        "What specific customer feedback, testimonials, or reviews highlight your unique advantages?",
        "What case studies, examples, or data points demonstrate your value proposition in action?",
    ], "Proof & Validation")

    # --- Section 7: UVP Synthesis ---
    add_section_header(doc, "7. UVP Synthesis")
    add_description(doc,
        "Synthesize the responses above into the deliverables below. Each differentiator should be "
        "named, specific, and connected to a customer pain point.")

    # Top 3 Differentiators
    add_styled_paragraph(doc, "Top 3 Differentiators:", font_size=11, bold=True, space_after=4)
    diff_table = doc.add_table(rows=3, cols=2)
    diff_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(3):
        row = diff_table.rows[i]
        set_cell_shading(row.cells[0], GREEN_HEX)
        set_cell_border(row.cells[0], top=BORDER_HEX, bottom=BORDER_HEX,
                       left=BORDER_HEX, right=BORDER_HEX)
        set_cell_border(row.cells[1], top=BORDER_HEX, bottom=BORDER_HEX,
                       left=BORDER_HEX, right=BORDER_HEX)
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(6.0)

        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(i + 1))
        run.font.name = "Inter"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        p2 = row.cells[1].paragraphs[0]
        run2 = p2.add_run("[Named differentiator + 2-3 sentence explanation]")
        run2.font.name = "Inter"
        run2.font.size = Pt(9.5)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # Draft UVP Statement
    add_styled_paragraph(doc, "Draft UVP Statement:", font_size=11, bold=True, space_after=4)
    uvp_table = doc.add_table(rows=1, cols=1)
    uvp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(uvp_table.rows[0].cells[0], LGGREEN_HEX)
    set_cell_border(uvp_table.rows[0].cells[0], top=GREEN_HEX, bottom=GREEN_HEX,
                   left=GREEN_HEX, right=GREEN_HEX)
    p = uvp_table.rows[0].cells[0].paragraphs[0]
    run = p.add_run("[2-3 sentences capturing core value proposition]")
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_paragraph()

    # Elevator Pitch
    add_styled_paragraph(doc, "Elevator Pitch (30 seconds):", font_size=11, bold=True, space_after=4)
    ep_table = doc.add_table(rows=1, cols=1)
    ep_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(ep_table.rows[0].cells[0], LGGREEN_HEX)
    set_cell_border(ep_table.rows[0].cells[0], top=GREEN_HEX, bottom=GREEN_HEX,
                   left=GREEN_HEX, right=GREEN_HEX)
    p = ep_table.rows[0].cells[0].paragraphs[0]
    run = p.add_run("[Natural-sounding paragraph: problem \u2192 audience \u2192 approach \u2192 proof]")
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_paragraph()

    # Positioning Statement
    add_styled_paragraph(doc, "Positioning Statement (Internal Use):", font_size=11, bold=True, space_after=4)
    ps_table = doc.add_table(rows=1, cols=1)
    ps_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(ps_table.rows[0].cells[0], LGGREEN_HEX)
    set_cell_border(ps_table.rows[0].cells[0], top=GREEN_HEX, bottom=GREEN_HEX,
                   left=GREEN_HEX, right=GREEN_HEX)
    p = ps_table.rows[0].cells[0].paragraphs[0]
    text = ("For [target audience]\n"
            "Who need [primary need]\n"
            "[Company] is the only [category]\n"
            "That [key differentiator]\n"
            "Because [proof/reason to believe].")
    run = p.add_run(text)
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "ClearLaunch_UVP_Template.docx")
    doc.save(out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    build_uvp_template()
