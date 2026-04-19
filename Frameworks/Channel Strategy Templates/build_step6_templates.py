"""
Build ClearLaunch Step 6 (Channel Strategy & Customer Journey) .docx template
with ClearThink brand styling. Industry-agnostic version.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Brand colors
GREEN = RGBColor(0x1B, 0x9B, 0x5E)
BRIGHT = RGBColor(0x3B, 0xEB, 0x96)
DARK = RGBColor(0x12, 0x17, 0x18)
GRAY = RGBColor(0x66, 0x66, 0x66)
PLACEHOLDER = RGBColor(0x99, 0x99, 0x99)

CREAM_HEX = "F6F3EF"
LGGREEN_HEX = "E5F5EC"
ALTROW_HEX = "F0F9F4"
GREEN_HEX = "1B9B5E"
DARK_HEX = "121718"
BORDER_HEX = "A8D9BD"
WHITE_HEX = "FFFFFF"


# ========================================
# REUSABLE HELPERS (from build_step4_step5)
# ========================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for side, color in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if color:
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            )
            borders.append(border)
    tcPr.append(borders)


def add_styled_paragraph(doc, text, font_name="Inter", font_size=11, bold=False,
                         color=DARK, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    return p


def add_section_header(doc, text):
    add_styled_paragraph(doc, text, font_size=14, bold=True, color=GREEN, space_after=4)


def add_subsection_header(doc, text):
    add_styled_paragraph(doc, text, font_size=12, bold=True, color=GREEN, space_after=4)


def add_description(doc, text):
    add_styled_paragraph(doc, text, font_size=9.5, color=GRAY, space_after=2)


def add_adaptive_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"Agent Note: {text}")
    run.font.name = "Inter"
    run.font.size = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = GREEN


def add_client_info_table(doc):
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_fields = ["Client / Business Name", "Industry / Vertical", "Date"]
    for i, label in enumerate(info_fields):
        row = table.rows[i]
        set_cell_shading(row.cells[0], LGGREEN_HEX)
        for cell in row.cells:
            set_cell_border(cell, top=BORDER_HEX, bottom=BORDER_HEX,
                           left=BORDER_HEX, right=BORDER_HEX)
        p = row.cells[0].paragraphs[0]
        run = p.add_run(label)
        run.font.name = "Inter"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = DARK

        p2 = row.cells[1].paragraphs[0]
        run2 = p2.add_run("[________________]")
        run2.font.name = "Inter"
        run2.font.size = Pt(10)
        run2.font.color.rgb = PLACEHOLDER
    doc.add_paragraph()


# ========================================
# NEW HELPERS FOR STEP 6
# ========================================

def add_data_table(doc, headers, rows, col_widths=None):
    """Create a branded data table with dark header and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for j, h in enumerate(headers):
        set_cell_shading(hdr.cells[j], DARK_HEX)
        set_cell_border(hdr.cells[j], top=BORDER_HEX, bottom=BORDER_HEX,
                       left=BORDER_HEX, right=BORDER_HEX)
        p = hdr.cells[j].paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Inter"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        bg = ALTROW_HEX if i % 2 == 0 else WHITE_HEX
        is_total = row_data[0].startswith("**") if row_data[0] else False
        if is_total:
            bg = LGGREEN_HEX

        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            set_cell_shading(cell, bg)
            set_cell_border(cell, top=BORDER_HEX, bottom=BORDER_HEX,
                           left=BORDER_HEX, right=BORDER_HEX)
            p = cell.paragraphs[0]
            display_text = cell_text.replace("**", "") if is_total else cell_text
            run = p.add_run(display_text)
            run.font.name = "Inter"
            run.font.size = Pt(9)
            run.font.bold = is_total
            run.font.color.rgb = DARK if not cell_text.startswith("[") else PLACEHOLDER

    # Apply column widths if provided
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)

    doc.add_paragraph()  # spacer
    return table


def add_callout_box(doc, text):
    """Create a single-cell callout box with light green background and green border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LGGREEN_HEX)
    set_cell_border(cell, top=GREEN_HEX, bottom=GREEN_HEX,
                   left=GREEN_HEX, right=GREEN_HEX)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Inter"
    run.font.size = Pt(9.5)
    run.font.color.rgb = PLACEHOLDER if text.startswith("[") else DARK
    doc.add_paragraph()


def add_bullet_list(doc, items, color=DARK):
    """Add a bulleted list with green triangle markers."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        # Green marker
        marker = p.add_run("\u25B8  ")
        marker.font.name = "Inter"
        marker.font.size = Pt(10)
        marker.font.color.rgb = GREEN
        # Text
        run = p.add_run(item)
        run.font.name = "Inter"
        run.font.size = Pt(10)
        run.font.color.rgb = color


def add_channel_role_card(doc, channel_name, role_title, tier, description):
    """Add a channel role description block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    # Channel name
    name_run = p.add_run(f"{channel_name} ")
    name_run.font.name = "Inter"
    name_run.font.size = Pt(11)
    name_run.font.bold = True
    name_run.font.color.rgb = DARK
    # Role title
    role_run = p.add_run(f"\u2014 {role_title} ")
    role_run.font.name = "Inter"
    role_run.font.size = Pt(11)
    role_run.font.italic = True
    role_run.font.color.rgb = GREEN
    # Tier badge
    tier_run = p.add_run(f"({tier})")
    tier_run.font.name = "Inter"
    tier_run.font.size = Pt(9)
    tier_run.font.bold = True
    tier_run.font.color.rgb = GREEN

    # Description
    add_styled_paragraph(doc, description, font_size=9.5, color=PLACEHOLDER, space_after=8)


# ========================================
# BUILD CHANNEL STRATEGY TEMPLATE
# ========================================

def build_channel_strategy_template():
    doc = Document()

    # ---- Title Block ----
    add_styled_paragraph(doc, "", space_after=12)
    add_styled_paragraph(doc, "CHANNEL STRATEGY & CUSTOMER JOURNEY",
                         font_size=24, bold=True, color=GREEN,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, "Strategy & Projections Template",
                         font_size=14, bold=False, color=DARK,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_paragraph(doc, "ClearLaunch System | Step 6: Channel Strategy & Customer Journey",
                         font_size=10, color=PLACEHOLDER,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    # Purpose
    add_styled_paragraph(
        doc,
        "Purpose: This template documents the complete channel strategy and customer journey for "
        "a ClearLaunch client. It evaluates all viable marketing channels, recommends a focused "
        "channel mix with budget allocation and performance projections, and maps the customer "
        "journey from awareness through purchase.",
        font_size=10, color=GRAY, space_after=12
    )

    # Client Info
    add_client_info_table(doc)

    # ================================================================
    # SECTION 1: EXECUTIVE SUMMARY
    # ================================================================
    add_section_header(doc, "1. Executive Summary")
    add_description(doc,
        "High-level overview of channel strategy findings. Populated after Phase A analysis.")

    add_callout_box(doc,
        "[Key Insight: 1-2 sentence summary of the strategic rationale for the recommended "
        "channel mix — why these channels, why this budget split, what the data shows.]")

    add_data_table(doc,
        headers=["Field", "Value"],
        rows=[
            ["Client", "[Client Name]"],
            ["Industry", "[Industry Context]"],
            ["Location", "[Geographic Market]"],
            ["Target Audience", "[ICP Description from Step 2]"],
            ["Recommended Channels", "[Primary 1], [Primary 2], [Supporting]"],
            ["Monthly Investment", "[Total monthly budget]"],
            ["Projected Monthly Revenue", "[Combined projected revenue]"],
        ],
        col_widths=[2.5, 4.0]
    )

    # ================================================================
    # SECTION 2: CUSTOMER JOURNEY MAP (Phase B)
    # ================================================================
    add_section_header(doc, "2. Customer Journey Map")
    add_description(doc,
        "Maps the customer experience from first awareness through purchase across the confirmed "
        "channels. This section is populated in Phase B after the client confirms the channel "
        "strategy.")
    add_adaptive_note(doc,
        "For B2B, journeys tend to be longer with more consideration-stage touchpoints. For local "
        "services, journeys are shorter with urgency-driven BOFU. For e-commerce, journeys may "
        "include browse/compare behaviors. For SaaS, free trial or demo often bridges MOFU to BOFU.")

    add_data_table(doc,
        headers=["Stage", "Customer Mindset", "Content / Touchpoints", "Channels", "Key KPI"],
        rows=[
            ["TOFU (Awareness)", "[What the customer is thinking/feeling]",
             "[Content types and touchpoints]", "[Active channels]", "[Primary metric]"],
            ["MOFU (Consideration)", "[What the customer is thinking/feeling]",
             "[Content types and touchpoints]", "[Active channels]", "[Primary metric]"],
            ["BOFU (Decision)", "[What the customer is thinking/feeling]",
             "[Content types and touchpoints]", "[Active channels]", "[Primary metric]"],
            ["Purchase", "[What the customer is thinking/feeling]",
             "[Content types and touchpoints]", "[Active channels]", "[Primary metric]"],
        ],
        col_widths=[1.1, 1.5, 1.5, 1.2, 1.2]
    )

    add_styled_paragraph(doc,
        "Journey Insights:",
        font_size=10, bold=True, space_after=4)
    add_styled_paragraph(doc,
        "[Key insights from the client decision call about how their customers currently find "
        "and engage with the business. Include any industry-specific journey patterns.]",
        font_size=9.5, color=PLACEHOLDER, space_after=8)

    # ================================================================
    # SECTION 3: CHANNEL UNIVERSE ASSESSMENT
    # ================================================================
    add_section_header(doc, "3. Channel Universe Assessment")
    add_description(doc,
        "All 7 channels in the ClearLaunch channel universe evaluated against the client's ICP, "
        "budget, and data availability.")

    add_data_table(doc,
        headers=["Channel", "Traffic Type", "Stage Served", "Data Avail.", "ICP Fit (1-5)", "Recommendation"],
        rows=[
            ["Google Ads", "Paid / Intent", "MOFU-BOFU", "High", "[score]", "[tier]"],
            ["Meta Ads", "Paid / Interruption", "TOFU-MOFU", "Moderate", "[score]", "[tier]"],
            ["LinkedIn Ads", "Paid / Professional", "TOFU-MOFU", "Moderate", "[score]", "[tier]"],
            ["TikTok", "Paid / Interruption", "TOFU", "Low", "[score]", "[tier]"],
            ["SEO", "Organic / Intent", "TOFU-BOFU", "High", "[score]", "[tier]"],
            ["Content Marketing", "Organic / Educational", "TOFU-MOFU", "Moderate", "[score]", "[tier]"],
            ["Organic Social", "Organic / Community", "TOFU", "Low", "[score]", "[tier]"],
        ],
        col_widths=[1.2, 1.1, 1.0, 0.8, 0.8, 1.2]
    )

    # Decision Rules
    add_subsection_header(doc, "3.1 Channel Selection Rationale")
    add_callout_box(doc,
        "[Decision Rules Applied:\n"
        "1. ICP Presence Gate \u2014 [which channels passed/failed and why]\n"
        "2. Budget Viability Gate \u2014 [which channels are affordable at this budget]\n"
        "3. Journey Coverage Check \u2014 [how the recommended set covers TOFU/MOFU/BOFU]\n"
        "4. Data Availability Tiebreaker \u2014 [any tiebreaker decisions made]]")

    # ================================================================
    # SECTION 4: RECOMMENDED CHANNEL STRATEGY
    # ================================================================
    add_section_header(doc, "4. Recommended Channel Strategy")
    add_description(doc,
        "Defines the role each recommended channel plays in the overall strategy.")

    add_subsection_header(doc, "4.1 Channel Roles")

    add_channel_role_card(doc,
        "[Primary Channel 1]", "[Role Title]", "PRIMARY",
        "[2-3 sentences explaining this channel's role, what funnel stages it serves, "
        "and why it's primary for this client.]")

    add_channel_role_card(doc,
        "[Primary Channel 2]", "[Role Title]", "PRIMARY",
        "[2-3 sentences explaining this channel's role, what funnel stages it serves, "
        "and why it's primary for this client.]")

    add_channel_role_card(doc,
        "[Supporting Channel]", "[Role Title]", "SUPPORTING",
        "[2-3 sentences explaining this channel's role, what funnel stages it serves, "
        "and why it's supporting rather than primary.]")

    # ================================================================
    # SECTION 5: INVESTMENT OVERVIEW
    # ================================================================
    add_section_header(doc, "5. Investment Overview")
    add_description(doc,
        "Budget allocation across recommended channels. Retainers are fixed management costs. "
        "Ad spend is the adjustable control variable.")

    add_data_table(doc,
        headers=["Channel", "Monthly Retainer", "Monthly Ad Spend", "Total Monthly", "Notes"],
        rows=[
            ["[Channel 1]", "[amount]", "[amount]", "[amount]", "[context]"],
            ["[Channel 2]", "[amount]", "[amount]", "[amount]", "[context]"],
            ["[Channel 3]", "[amount]", "[amount]", "[amount]", "[context]"],
            ["**Total", "[total retainers]", "[total ad spend]", "[total monthly]", ""],
        ],
        col_widths=[1.3, 1.2, 1.2, 1.2, 1.3]
    )

    add_callout_box(doc,
        "[Key Insight: Note about budget philosophy \u2014 retainers are fixed operational costs, "
        "ad spend is the adjustable control variable. Scaling up means increasing ad spend, not "
        "adding channels.]")

    # ================================================================
    # SECTION 6: CHANNEL PROJECTIONS — DATA-BACKED
    # ================================================================
    add_section_header(doc, "6. Channel Projections \u2014 Data-Backed")
    add_description(doc,
        "Projections built from real keyword data (Step 3 Market Research). Bottom-up methodology: "
        "Clicks = Budget \u00f7 CPC \u2192 Leads \u2192 Clients \u2192 Revenue.")

    add_subsection_header(doc, "6.1 [Data-Backed Channel] \u2014 Service Line Comparison")

    add_data_table(doc,
        headers=["Metric", "[Service Line 1]", "[Service Line 2]", "Notes"],
        rows=[
            ["DFY Management Retainer", "[amount]", "\u2014", "Fixed monthly"],
            ["Avg CPC", "[amount]", "[amount]", "From Step 3 keyword data"],
            ["Monthly Ad Budget", "[amount]", "[amount]", "Input"],
            ["LP Conversion Rate", "[rate]", "[rate]", "Input"],
            ["Sales Closing Rate", "[rate]", "[rate]", "Input"],
            ["Revenue Per Client", "[amount]", "[amount]", "From Step 5"],
            ["", "", "", ""],
            ["Monthly Clicks", "[calculated]", "[calculated]", "= Budget \u00f7 CPC"],
            ["Monthly Leads", "[calculated]", "[calculated]", "= Clicks \u00d7 Conv Rate"],
            ["Cost Per Lead", "[calculated]", "[calculated]", "= Budget \u00f7 Leads"],
            ["New Clients/Month", "[calculated]", "[calculated]", "= Leads \u00d7 Close Rate"],
            ["Cost Per Acquisition", "[calculated]", "[calculated]", "= Budget \u00f7 Clients"],
            ["Monthly Revenue", "[calculated]", "[calculated]", "= Clients \u00d7 Rev/Client"],
            ["Ad Spend ROAS", "[calculated]", "[calculated]", "= Revenue \u00f7 Ad Spend"],
            ["Monthly Profit", "[calculated]", "[calculated]", "= Revenue \u2212 Ad Spend"],
            ["Profit Margin", "[calculated]", "[calculated]", "= Profit \u00f7 Revenue"],
        ],
        col_widths=[1.5, 1.3, 1.3, 2.0]
    )

    # Combined profitability
    add_styled_paragraph(doc, "Combined Profitability (includes retainer):",
                         font_size=10, bold=True, space_after=4)
    add_data_table(doc,
        headers=["Metric", "Combined"],
        rows=[
            ["Total Revenue", "[sum of service lines]"],
            ["Total Ad Spend", "[sum of service line budgets]"],
            ["Total Investment", "[ad spend + retainer]"],
            ["New Clients", "[sum]"],
            ["True ROAS", "[Revenue \u00f7 Total Investment]"],
            ["CPA", "[Total Investment \u00f7 Clients]"],
            ["Profit", "[Revenue \u2212 Total Investment]"],
            ["Profit Margin", "[Profit \u00f7 Revenue]"],
        ],
        col_widths=[3.0, 3.0]
    )

    # ================================================================
    # SECTION 7: CHANNEL PROJECTIONS — MODELED
    # ================================================================
    add_section_header(doc, "7. Channel Projections \u2014 Modeled Estimates")
    add_description(doc,
        "Projections built from industry CPM/CTR benchmarks. Top-down methodology: "
        "Impressions = Budget \u00f7 CPM \u00d7 1000 \u2192 Clicks \u2192 Leads \u2192 Clients \u2192 Revenue.")

    add_subsection_header(doc, "7.1 [Modeled Channel] \u2014 Service Line Comparison")

    add_data_table(doc,
        headers=["Metric", "[Service Line 1]", "[Service Line 2]", "Notes"],
        rows=[
            ["DFY Management Retainer", "[amount]", "\u2014", "Fixed monthly"],
            ["Daily Ad Budget", "[amount]", "[amount]", "Input"],
            ["Monthly Ad Budget", "[calculated]", "[calculated]", "= Daily \u00d7 30"],
            ["CPM", "[amount]", "[amount]", "Industry benchmark"],
            ["CTR", "[rate]", "[rate]", "Industry benchmark"],
            ["Effective CPC", "[calculated]", "[calculated]", "= CPM \u00f7 (CTR \u00d7 1000)"],
            ["LP Conversion Rate", "[rate]", "[rate]", "Input"],
            ["Sales Closing Rate", "[rate]", "[rate]", "Input"],
            ["Revenue Per Client", "[amount]", "[amount]", "From Step 5"],
            ["", "", "", ""],
            ["Monthly Impressions", "[calculated]", "[calculated]", "= Budget \u00f7 CPM \u00d7 1000"],
            ["Monthly Clicks", "[calculated]", "[calculated]", "= Impressions \u00d7 CTR"],
            ["Monthly Leads", "[calculated]", "[calculated]", "= Clicks \u00d7 Conv Rate"],
            ["Cost Per Lead", "[calculated]", "[calculated]", "= Budget \u00f7 Leads"],
            ["New Clients/Month", "[calculated]", "[calculated]", "= Leads \u00d7 Close Rate"],
            ["Cost Per Acquisition", "[calculated]", "[calculated]", "= Budget \u00f7 Clients"],
            ["Monthly Revenue", "[calculated]", "[calculated]", "= Clients \u00d7 Rev/Client"],
            ["Ad Spend ROAS", "[calculated]", "[calculated]", "= Revenue \u00f7 Ad Spend"],
            ["Monthly Profit", "[calculated]", "[calculated]", "= Revenue \u2212 Ad Spend"],
            ["Profit Margin", "[calculated]", "[calculated]", "= Profit \u00f7 Revenue"],
        ],
        col_widths=[1.5, 1.3, 1.3, 2.0]
    )

    # Benchmark sources
    add_styled_paragraph(doc, "Benchmark Sources:",
                         font_size=10, bold=True, space_after=4)
    add_data_table(doc,
        headers=["Source", "Metric", "Value", "Notes"],
        rows=[
            ["[Source 1]", "CPM", "[range]", "[context]"],
            ["[Source 2]", "CTR", "[range]", "[context]"],
            ["[Source 3]", "LP Conv Rate", "[range]", "[context]"],
            ["[Source 4]", "Close Rate", "[range]", "[context]"],
        ],
        col_widths=[1.8, 1.2, 1.2, 2.0]
    )

    # Combined profitability
    add_styled_paragraph(doc, "Combined Profitability (includes retainer):",
                         font_size=10, bold=True, space_after=4)
    add_data_table(doc,
        headers=["Metric", "Combined"],
        rows=[
            ["Total Revenue", "[sum of service lines]"],
            ["Total Ad Spend", "[sum of service line budgets]"],
            ["Total Investment", "[ad spend + retainer]"],
            ["New Clients", "[sum]"],
            ["True ROAS", "[Revenue \u00f7 Total Investment]"],
            ["CPA", "[Total Investment \u00f7 Clients]"],
            ["Profit", "[Revenue \u2212 Total Investment]"],
            ["Profit Margin", "[Profit \u00f7 Revenue]"],
        ],
        col_widths=[3.0, 3.0]
    )

    # ================================================================
    # SECTION 8: CONTENT & SEO FOUNDATION
    # ================================================================
    add_section_header(doc, "8. Content & SEO Foundation")
    add_description(doc,
        "Inventory of foundation pages and SEO services needed to support the channel strategy. "
        "Content is the storefront \u2014 channels drive traffic to it.")

    add_subsection_header(doc, "8.1 Foundation Page Inventory")

    add_data_table(doc,
        headers=["Deliverable", "Pages", "Asset Value", "Notes"],
        rows=[
            ["[Page Category 1]", "[qty]", "[cost]", "[description]"],
            ["[Page Category 2]", "[qty]", "[cost]", "[description]"],
            ["[Page Category 3]", "[qty]", "[cost]", "[description]"],
            ["[Page Category 4]", "[qty]", "[cost]", "[description]"],
            ["[Page Category 5]", "[qty]", "[cost]", "[description]"],
            ["[Page Category 6]", "[qty]", "[cost]", "[description]"],
            ["**Content Marketing Subtotal", "[total pages]", "[total cost]", ""],
        ],
        col_widths=[2.0, 0.8, 1.2, 2.2]
    )

    add_subsection_header(doc, "8.2 SEO Services")

    add_data_table(doc,
        headers=["Service", "Details", "Monthly Cost"],
        rows=[
            ["[Service 1]", "[description]", "[cost or range]"],
            ["[Service 2]", "[description]", "[cost or range]"],
            ["[Service 3]", "[description]", "[cost or range]"],
        ],
        col_widths=[2.0, 2.5, 1.5]
    )

    add_callout_box(doc,
        "[Key Insight: Content is the foundation that all channels drive traffic to. Without "
        "landing pages, blog content, and a strong SEO base, paid channels have nowhere to send "
        "prospects. Content marketing is a one-time investment that compounds over time.]")

    # ================================================================
    # SECTION 9: KEY TAKEAWAYS & NEXT STEPS
    # ================================================================
    add_section_header(doc, "9. Key Takeaways & Next Steps")
    add_description(doc,
        "Top findings and bridge to Step 7 (Success Metrics & Launch Roadmap).")

    add_subsection_header(doc, "Key Takeaways")

    add_bullet_list(doc, [
        "[Takeaway about primary channel role and why it's the foundation]",
        "[Takeaway about secondary channel role and how it complements]",
        "[Takeaway about supporting channel role and demand generation]",
        "[Takeaway about funnel architecture and how channels work together]",
        "[Takeaway about budget control lever and scaling strategy]",
    ], color=PLACEHOLDER)

    add_subsection_header(doc, "Next Steps \u2014 Bridge to Step 7")

    add_callout_box(doc,
        "Step 6 answered: \"Where should we show up, and what should we expect?\"\n\n"
        "Step 7 will answer: \"How do we execute, in what order, and how do we know it's working?\"\n\n"
        "\u2022  Launch sequencing and 90-day roadmap (Step 7)\n"
        "\u2022  KPI targets with success/warning/failure thresholds (Step 7)\n"
        "\u2022  Weekly milestone tracking framework (Step 7)")

    # ---- Save ----
    doc.save("Frameworks/ClearLaunch_Step6_ChannelStrategy_Template.docx")
    print("Created: ClearLaunch_Step6_ChannelStrategy_Template.docx")


if __name__ == "__main__":
    build_channel_strategy_template()
    print("\nDone! Template created in Frameworks/")
